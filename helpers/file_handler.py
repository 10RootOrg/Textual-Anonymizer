# helpers/file_handler.py
import os
import json5
import sys
import logging
import mimetypes
from pathlib import Path

# Text/CSV handling
import csv
import io

# Excel handling
import pandas as pd
import openpyxl

# Word document handling
import docx

# PDF handling
import PyPDF2
from pdfminer.high_level import extract_text as pdf_extract_text

# JSON/XML handling
import xml.etree.ElementTree as ET

def load_config(logger=None):
    """
    Load the JSON5 configuration file from the same directory as the script
    
    Args:
        logger (logging.Logger, optional): Logger instance
        
    Returns:
        dict: Configuration dictionary
    """
    # Get the directory where the script is located
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(script_dir, 'config.json5')
    
    if logger:
        logger.debug(f"Loading config from: {config_path}")
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json5.load(f)
            
        # Log configuration details if logger is provided
        if logger:
            keyword_count = len(config.get('keywords', []))
            censor_type = config.get('censor_type', 'regex')
            logger.debug(f"Config loaded: {keyword_count} keywords, type: {censor_type}")
            
        return config
    except FileNotFoundError:
        error_msg = f"Config file not found at {config_path}"
        if logger:
            logger.error(error_msg)
        raise FileNotFoundError(error_msg)
    except json5.decode.JSONDecodeError as e:
        error_msg = f"Invalid JSON5 format in config file: {str(e)}"
        if logger:
            logger.error(error_msg)
        raise ValueError(error_msg)
    except Exception as e:
        error_msg = f"Error loading config: {str(e)}"
        if logger:
            logger.error(error_msg)
        raise Exception(error_msg)

def load_file(file_path, logger):
    """
    Load a file's content regardless of its format
    
    Args:
        file_path (str): Path to the file to load
        logger (logging.Logger): Logger instance
        
    Returns:
        tuple: (content, format, metadata)
            - content: The file content as text or structured data
            - format: String indicating the file format
            - metadata: Dictionary with additional file information
            
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file format is unsupported
        Exception: For other errors during loading
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not os.path.isfile(file_path):
        logger.error(f"Not a file: {file_path}")
        raise ValueError(f"Not a file: {file_path}")
    
    # Determine file type
    file_ext = os.path.splitext(file_path)[1].lower()
    mime_type, _ = mimetypes.guess_type(file_path)
    
    logger.info(f"Loading file: {file_path}")
    logger.debug(f"File extension: {file_ext}, MIME type: {mime_type}")
    
    metadata = {
        'path': file_path,
        'size': os.path.getsize(file_path),
        'extension': file_ext,
        'mime_type': mime_type
    }
    
    try:
        # Handle different file types
        # Plain text files
        if file_ext in ['.txt', '.log', '.md', '.rst', '.csv', '.tsv']:
            return load_text_file(file_path, file_ext, logger, metadata)
            
        # Excel files
        elif file_ext in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.odf', '.ods', '.odt']:
            return load_excel_file(file_path, logger, metadata)
            
        # Word documents
        elif file_ext in ['.doc', '.docx']:
            return load_word_file(file_path, logger, metadata)
            
        # PDF files
        elif file_ext == '.pdf':
            return load_pdf_file(file_path, logger, metadata)
            
        # JSON and JSON5 files
        elif file_ext in ['.json', '.json5']:
            return load_json_file(file_path, file_ext, logger, metadata)
            
        # XML files
        elif file_ext in ['.xml', '.html', '.htm']:
            return load_xml_file(file_path, logger, metadata)
            
        # Binary files - return empty content but metadata
        elif file_ext in ['.zip', '.rar', '.tar', '.gz', '.7z', '.exe', '.bin', '.dat']:
            logger.warning(f"Binary file format not supported for content extraction: {file_ext}")
            return "", "binary", metadata
            
        # Unknown format
        else:
            # Try to load as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                logger.info(f"Loaded unknown file as text: {file_path}")
                return content, "text", metadata
            except UnicodeDecodeError:
                # If not text, inform it's a binary file
                logger.warning(f"Unsupported file format: {file_ext}")
                return "", "unsupported", metadata
                
    except Exception as e:
        logger.error(f"Error loading file {file_path}: {str(e)}", exc_info=True)
        raise

def load_text_file(file_path, file_ext, logger, metadata):
    """Load a text-based file"""
    format_name = "text"
    
    try:
        # Handle CSV/TSV specially
        if file_ext in ['.csv', '.tsv']:
            delimiter = ',' if file_ext == '.csv' else '\t'
            
            # Read as structured data
            with open(file_path, 'r', encoding='utf-8', newline='') as f:
                reader = csv.reader(f, delimiter=delimiter)
                rows = list(reader)
                
            # Also get as plain text for processing
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            metadata['rows'] = len(rows)
            metadata['columns'] = len(rows[0]) if rows else 0
            format_name = "csv" if file_ext == '.csv' else "tsv"
            
            logger.info(f"Loaded {format_name} file: {len(rows)} rows")
            return content, format_name, metadata
        
        # Regular text files
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        lines = content.count('\n') + 1
        metadata['lines'] = lines
        logger.info(f"Loaded text file: {lines} lines")
        return content, format_name, metadata
        
    except UnicodeDecodeError:
        # Try with different encoding
        logger.warning("UTF-8 decoding failed, trying with ISO-8859-1 encoding")
        with open(file_path, 'r', encoding='ISO-8859-1') as f:
            content = f.read()
        lines = content.count('\n') + 1
        metadata['lines'] = lines
        metadata['encoding'] = 'ISO-8859-1'
        logger.info(f"Loaded text file with ISO-8859-1 encoding: {lines} lines")
        return content, format_name, metadata

def load_excel_file(file_path, logger, metadata):
    """Load an Excel file"""
    # Load the Excel file
    try:
        # Get sheet names
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames
        metadata['sheets'] = sheet_names
        wb.close()
        
        # Read content
        content = ""
        sheet_data = {}
        
        for sheet in sheet_names:
            # Read sheet into DataFrame
            df = pd.read_excel(file_path, sheet_name=sheet)
            
            # Convert to string representation
            string_rep = df.to_string(index=False)
            content += f"--- Sheet: {sheet} ---\n{string_rep}\n\n"
            
            # Store structured data
            sheet_data[sheet] = df.to_dict('records')
        
        metadata['row_counts'] = {sheet: len(sheet_data[sheet]) for sheet in sheet_data}
        metadata['data'] = sheet_data  # Store structured data in metadata
        
        logger.info(f"Loaded Excel file: {len(sheet_names)} sheets")
        return content, "excel", metadata
        
    except Exception as e:
        logger.error(f"Error loading Excel file: {str(e)}")
        raise

def load_word_file(file_path, logger, metadata):
    """Load a Word document"""
    try:
        # Only .docx is supported by python-docx
        if file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text for cell in row.cells])
                    content += f"\n{row_text}"
            
            metadata['paragraphs'] = len(doc.paragraphs)
            metadata['tables'] = len(doc.tables)
            
            logger.info(f"Loaded Word document: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            return content, "docx", metadata
        else:
            logger.error("Only .docx format is supported for Word documents")
            raise ValueError("Only .docx format is supported for Word documents")
            
    except Exception as e:
        logger.error(f"Error loading Word document: {str(e)}")
        raise

def load_pdf_file(file_path, logger, metadata):
    """Load a PDF file"""
    content = ""
    
    # Try with PyPDF2 first
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfFileReader(file)
            metadata['pages'] = reader.getNumPages()
            
            # Extract text from each page
            for page_num in range(reader.getNumPages()):
                page = reader.getPage(page_num)
                content += page.extractText() + "\n"
                
            logger.info(f"Loaded PDF file with PyPDF2: {reader.getNumPages()} pages")
            
            # If content is very sparse, try pdfminer for better extraction
            if len(content.strip()) < 100 and reader.getNumPages() > 0:
                logger.info("PyPDF2 extraction produced limited content, trying pdfminer...")
                raise Exception("Sparse content, trying alternative method")
                
            return content, "pdf", metadata
            
    except Exception as e:
        logger.debug(f"PyPDF2 extraction issue: {str(e)}")
    
    # Fall back to pdfminer
    try:
        content = pdf_extract_text(file_path)
        # Estimate number of pages (rough estimate based on newlines)
        pages = max(1, content.count('\f') + 1)
        metadata['pages'] = pages
        
        logger.info(f"Loaded PDF file with pdfminer: approximately {pages} pages")
        return content, "pdf", metadata
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
    
    raise Exception("Failed to extract text from PDF file")

def load_json_file(file_path, file_ext, logger, metadata):
    """Load a JSON or JSON5 file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            if file_ext == '.json5':
                data = json5.load(f)
            else:
                import json
                data = json.load(f)
        
        # Convert to string representation for text processing
        if file_ext == '.json5':
            content = json5.dumps(data, indent=2)
        else:
            import json
            content = json.dumps(data, indent=2)
            
        metadata['data'] = data  # Store structured data in metadata
        
        logger.info(f"Loaded JSON file successfully")
        return content, "json", metadata
        
    except Exception as e:
        logger.error(f"Error loading JSON file: {str(e)}")
        raise

def load_xml_file(file_path, logger, metadata):
    """Load an XML file"""
    try:
        # Parse XML
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # Convert to string representation
        content = ET.tostring(root, encoding='unicode', method='xml')
        
        metadata['root_tag'] = root.tag
        
        logger.info(f"Loaded XML file successfully")
        return content, "xml", metadata
        
    except Exception as e:
        logger.error(f"Error loading XML file: {str(e)}")
        raise

def save_file(content, output_path, format_type, original_metadata, logger):
    """
    Save content to a file based on the specified format
    
    Args:
        content (str or object): Content to save
        output_path (str): Path where to save the file
        format_type (str): Format type of the content
        original_metadata (dict): Metadata from the original file
        logger (logging.Logger): Logger instance
        
    Returns:
        bool: True if successful, False otherwise
        
    Raises:
        ValueError: If the format is unsupported for saving
        Exception: For other errors during saving
    """
    try:
        logger.info(f"Saving file to: {output_path}")
        
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            logger.debug(f"Created directory: {output_dir}")
        
        # Determine output format based on extension
        output_ext = os.path.splitext(output_path)[1].lower()
        
        # If no extension, use the original format
        if not output_ext and 'extension' in original_metadata:
            output_ext = original_metadata['extension']
            output_path += output_ext
            logger.debug(f"Added extension to output path: {output_path}")
        
        # Handle different file types for saving
        if output_ext in ['.txt', '.log', '.md', '.rst']:
            # Simple text file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Saved text file: {output_path}")
            
        elif output_ext in ['.csv', '.tsv']:
            # CSV/TSV file
            delimiter = ',' if output_ext == '.csv' else '\t'
            
            # If content is already structured (list of lists)
            if isinstance(content, list) and all(isinstance(row, list) for row in content):
                with open(output_path, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f, delimiter=delimiter)
                    writer.writerows(content)
            else:
                # Just write as text
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            logger.info(f"Saved {'CSV' if output_ext == '.csv' else 'TSV'} file: {output_path}")
            
        elif output_ext in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.odf', '.ods', '.odt']:
            # Excel file
            # Check if we have structured data in metadata
            if 'data' in original_metadata and isinstance(original_metadata['data'], dict):
                # Create Excel writer
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    # Write each sheet
                    for sheet_name, data in original_metadata['data'].items():
                        df = pd.DataFrame(data)
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
            else:
                # Create a simple one-sheet Excel with the content
                df = pd.DataFrame([content.split('\n')])
                df.to_excel(output_path, index=False, header=False)
                
            logger.info(f"Saved Excel file: {output_path}")
            
        elif output_ext in ['.docx']:
            # Word document
            # Create a new document
            doc = docx.Document()
            
            # Add content as paragraphs
            for paragraph in content.split('\n'):
                doc.add_paragraph(paragraph)
                
            # Save the document
            doc.save(output_path)
            logger.info(f"Saved Word document: {output_path}")
            
        elif output_ext == '.pdf':
            logger.error("Direct PDF writing not supported. PDF requires specialized libraries for creation.")
            raise ValueError("Direct PDF writing not supported")
            
        elif output_ext in ['.json', '.json5']:
            # JSON file
            if isinstance(content, str):
                # Try to parse string as JSON
                try:
                    import json
                    data = json.loads(content)
                except json.JSONDecodeError:
                    # If not valid JSON, write as text
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    logger.info(f"Saved content as text to JSON file: {output_path}")
                    return True
            else:
                # Use the content directly
                data = content
                
            # Write JSON file
            with open(output_path, 'w', encoding='utf-8') as f:
                if output_ext == '.json5':
                    json5.dump(data, f, indent=2)
                else:
                    import json
                    json.dump(data, f, indent=2)
                    
            logger.info(f"Saved JSON file: {output_path}")
            
        elif output_ext in ['.xml', '.html', '.htm']:
            # XML file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Saved XML/HTML file: {output_path}")
            
        else:
            # Default: save as text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Saved file as text: {output_path}")
            
        return True
        
    except Exception as e:
        logger.error(f"Error saving file to {output_path}: {str(e)}", exc_info=True)
        raise

def process_file(input_path, output_path, config, logger):
    """
    Process a file by loading, censoring, and saving its content
    """
    try:
        # Import censoring functions
        from censor.regex_censor import censor_text_regex
        from censor.offline_model_censor import censor_text_presidio
        
        # Get censoring configuration
        censor_types = config.get('censor_type', 'regex')
        keywords = config.get('keywords', [])
        model_name = config.get('censor_model', "")
        
        # Convert to list if it's a string
        if isinstance(censor_types, str):
            censor_types = [censor_types]
        
        # Get file extension
        file_ext = os.path.splitext(input_path)[1].lower()
        
        # Handle complex document formats specially
        if file_ext == '.docx':
            # For DOCX files, process with format preservation
            return process_docx_with_preserving_format(input_path, output_path, censor_types, keywords, model_name, logger)
        elif file_ext == '.pdf':
            # For PDF files, process with format preservation
            return process_pdf_with_preserving_format(input_path, output_path, censor_types, keywords, model_name, logger)
        else:
            # For other file types, use the existing approach
            content, format_type, metadata = load_file(input_path, logger)
            
            # Apply each censoring method in sequence
            censored_content = content
            for censor_type in censor_types:
                censor_type = censor_type.lower()
                logger.info(f"Applying censoring method: {censor_type}")
                
                if censor_type == 'offline_model':
                    censored_content = censor_text_presidio(censored_content, keywords, model_name, logger)
                elif censor_type == 'regex':
                    censored_content = censor_text_regex(censored_content, keywords, logger)
                else:
                    logger.warning(f"Unknown censoring method: {censor_type}")
            
            # Save the censored content
            logger.info("Saving censored content")
            save_file(censored_content, output_path, format_type, metadata, logger)
        
        logger.info("File processed successfully")
        return True
        
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}", exc_info=True)
        raise

def process_docx_with_preserving_format(input_path, output_path, censor_types, keywords, model_name, logger):
    """
    Process a DOCX file by censoring text while preserving formatting
    """
    import docx
    
    logger.info(f"Processing DOCX file with format preservation: {input_path}")
    
    # Load the document
    doc = docx.Document(input_path)
    
    # Initialize models ONCE before processing
    regex_censoring = 'regex' in censor_types
    presidio_censoring = 'offline_model' in censor_types
    
    # If using Presidio, initialize it once with loaded model
    analyzer = None
    anonymizer = None
    if presidio_censoring:
        # Import and initialize Presidio only once
        from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
        from presidio_analyzer.nlp_engine import SpacyNlpEngine
        from presidio_anonymizer import AnonymizerEngine
        from presidio_anonymizer.entities import OperatorConfig
        from censor.offline_model_censor import ensure_model_exists
        
        # Load model once
        model_path = ensure_model_exists(model_name, logger)
        logger.info(f"Creating NLP engine with model: {model_path}")
        
        # Create NLP engine with the model
        nlp_engine = SpacyNlpEngine(models=[{"lang_code": "en", "model_name": model_path}])
        
        # Create registry for recognizers
        registry = RecognizerRegistry()
        
        # Create analyzer and anonymizer engines
        analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
        anonymizer = AnonymizerEngine()
    
    # Define efficient censoring function that uses pre-loaded models
    def efficient_censoring(text):
        result = text
        
        # Apply regex if needed
        if regex_censoring:
            from censor.regex_censor import censor_text_regex
            result = censor_text_regex(result, keywords, logger)
        
        # Apply Presidio if needed (with pre-loaded model)
        if presidio_censoring and analyzer and anonymizer:
            # Define operators for anonymization - include all the entities you want to censor
            operators = {
                "PERSON": OperatorConfig("replace", {"new_value": "[PERSON_NAME]"}),
                "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL]"}),
                "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[PHONE]"}),
                "CREDIT_CARD": OperatorConfig("replace", {"new_value": "[CREDIT_CARD]"}),
                "DOMAIN": OperatorConfig("replace", {"new_value": "[DOMAIN]"}),
                "DATE_TIME": OperatorConfig("replace", {"new_value": "[DATE]"}),
                "LOCATION": OperatorConfig("replace", {"new_value": "[ADDRESS]"}),
                "NRP": OperatorConfig("replace", {"new_value": "[ID_NUMBER]"}),
                "IP_ADDRESS": OperatorConfig("replace", {"new_value": "[IP_ADDRESS]"}),
                "US_SSN": OperatorConfig("replace", {"new_value": "[SSN_ID]"}),
                "US_BANK_NUMBER": OperatorConfig("replace", {"new_value": "[BANK_ACCOUNT]"}),
                "IBAN_CODE": OperatorConfig("replace", {"new_value": "[BANK_ACCOUNT]"}),
                "US_DRIVER_LICENSE": OperatorConfig("replace", {"new_value": "[DRIVER_LICENSE]"}),
                "CRYPTO": OperatorConfig("replace", {"new_value": "[CRYPTO_ADDRESS]"}),
                "URL": OperatorConfig("replace", {"new_value": "[URL]"}),
                "COMPANY": OperatorConfig("replace", {"new_value": "[COMPANY_NAME]"}),
                "MONEY": OperatorConfig("replace", {"new_value": "[MONEY]"}),
                "PASSWORD": OperatorConfig("replace", {"new_value": "[PASSWORD]"}),
                "API_KEY": OperatorConfig("replace", {"new_value": "[API_KEY]"})
            }
            
            # Custom simplified version of censor_text_presidio that doesn't reload model
            if result:
                # Analyze the text
                analyzer_results = analyzer.analyze(text=result, language="en")
                
                # Anonymize if results found
                if analyzer_results:
                    anonymizer_result = anonymizer.anonymize(
                        text=result,
                        analyzer_results=analyzer_results,
                        operators=operators
                    )
                    result = anonymizer_result.text
        
        return result
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            original_text = paragraph.text
            censored_text = efficient_censoring(original_text)
            
            if original_text != censored_text:
                # Clear paragraph content
                for run in paragraph.runs:
                    run.text = ""
                
                # Add censored content to first run
                if paragraph.runs:
                    paragraph.runs[0].text = censored_text
                else:
                    paragraph.add_run(censored_text)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        censored_text = efficient_censoring(original_text)
                        
                        if original_text != censored_text:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = censored_text
                            else:
                                paragraph.add_run(censored_text)
    
    # Process headers and footers
    for section in doc.sections:
        # Process header
        if section.header:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    censored_text = efficient_censoring(original_text)
                    
                    if original_text != censored_text:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = censored_text
                        else:
                            paragraph.add_run(censored_text)
        
        # Process footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    censored_text = efficient_censoring(original_text)
                    
                    if original_text != censored_text:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = censored_text
                        else:
                            paragraph.add_run(censored_text)
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Saved censored DOCX with preserved formatting: {output_path}")
    
    return True

def process_pdf_with_preserving_format(input_path, output_path, censor_types, keywords, model_name, logger):
    """
    Process a PDF file by censoring text while preserving layout and images
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF not installed. Install with: pip install pymupdf")
        raise ImportError("PyMuPDF required for PDF processing with format preservation")
    
    logger.info(f"Processing PDF file with format preservation: {input_path}")
    
    # Initialize models ONCE before processing
    regex_censoring = 'regex' in censor_types
    presidio_censoring = 'offline_model' in censor_types
    
    # If using Presidio, initialize it once with loaded model
    analyzer = None
    anonymizer = None
    if presidio_censoring:
        # Import and initialize Presidio only once
        from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
        from presidio_analyzer.nlp_engine import SpacyNlpEngine
        from presidio_anonymizer import AnonymizerEngine
        from presidio_anonymizer.entities import OperatorConfig
        from censor.offline_model_censor import ensure_model_exists
        
        # Load model once
        model_path = ensure_model_exists(model_name, logger)
        logger.info(f"Creating NLP engine with model: {model_path}")
        
        # Create NLP engine with the model
        nlp_engine = SpacyNlpEngine(models=[{"lang_code": "en", "model_name": model_path}])
        
        # Create registry for recognizers
        registry = RecognizerRegistry()
        
        # Create analyzer and anonymizer engines
        analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
        anonymizer = AnonymizerEngine()
    
    # Define efficient censoring function that uses pre-loaded models
    def efficient_censoring(text):
        result = text
        
        # Apply regex if needed
        if regex_censoring:
            from censor.regex_censor import censor_text_regex
            result = censor_text_regex(result, keywords, logger)
        
        # Apply Presidio if needed (with pre-loaded model)
        if presidio_censoring and analyzer and anonymizer:
            # Define operators for anonymization - include all the entities you want to censor
            operators = {
                "PERSON": OperatorConfig("replace", {"new_value": "[PERSON_NAME]"}),
                "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL]"}),
                "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[PHONE]"}),
                "CREDIT_CARD": OperatorConfig("replace", {"new_value": "[CREDIT_CARD]"}),
                "DOMAIN": OperatorConfig("replace", {"new_value": "[DOMAIN]"}),
                "DATE_TIME": OperatorConfig("replace", {"new_value": "[DATE]"}),
                "LOCATION": OperatorConfig("replace", {"new_value": "[ADDRESS]"}),
                "NRP": OperatorConfig("replace", {"new_value": "[ID_NUMBER]"}),
                "IP_ADDRESS": OperatorConfig("replace", {"new_value": "[IP_ADDRESS]"}),
                "US_SSN": OperatorConfig("replace", {"new_value": "[SSN_ID]"}),
                "US_BANK_NUMBER": OperatorConfig("replace", {"new_value": "[BANK_ACCOUNT]"}),
                "IBAN_CODE": OperatorConfig("replace", {"new_value": "[BANK_ACCOUNT]"}),
                "US_DRIVER_LICENSE": OperatorConfig("replace", {"new_value": "[DRIVER_LICENSE]"}),
                "CRYPTO": OperatorConfig("replace", {"new_value": "[CRYPTO_ADDRESS]"}),
                "URL": OperatorConfig("replace", {"new_value": "[URL]"}),
                "COMPANY": OperatorConfig("replace", {"new_value": "[COMPANY_NAME]"}),
                "MONEY": OperatorConfig("replace", {"new_value": "[MONEY]"}),
                "PASSWORD": OperatorConfig("replace", {"new_value": "[PASSWORD]"}),
                "API_KEY": OperatorConfig("replace", {"new_value": "[API_KEY]"})
            }
            
            # Custom simplified version of censor_text_presidio that doesn't reload model
            if result:
                # Analyze the text
                analyzer_results = analyzer.analyze(text=result, language="en")
                
                # Anonymize if results found
                if analyzer_results:
                    anonymizer_result = anonymizer.anonymize(
                        text=result,
                        analyzer_results=analyzer_results,
                        operators=operators
                    )
                    result = anonymizer_result.text
        
        return result
    
    # Open the PDF
    doc = fitz.open(input_path)
    
    # Process each page
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get text blocks
        text_blocks = page.get_text("blocks")
        
        for block in text_blocks:
            if len(block) >= 5 and block[4].strip():  # If there's text (index 4 contains text)
                original_text = block[4]
                censored_text = efficient_censoring(original_text)
                
                if original_text != censored_text:
                    # Create rectangle covering original text
                    rect = fitz.Rect(block[0], block[1], block[2], block[3])
                    
                    # Add redaction annotation
                    annot = page.add_redact_annot(rect, fill=(1, 1, 1))  # White fill
                    
                    # Apply redaction to remove original text
                    page.apply_redactions()
                    
                    # Insert the new text
                    font_size = 11  # Estimate font size
                    page.insert_text(
                        (block[0], block[1] + font_size),
                        censored_text,
                        fontsize=font_size
                    )
    
    # Save the modified PDF
    doc.save(output_path)
    doc.close()
    
    logger.info(f"Saved censored PDF with preserved layout: {output_path}")
    return True